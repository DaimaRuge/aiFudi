import os
import tempfile
from typing import Optional, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document as DocxDocument
import logging
from unstructured.partition.auto import partition
from unstructured.cleaners.core import clean_extra_whitespace

from models import DocumentType

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器，支持多种格式"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,
            '.txt': DocumentType.TXT,
            '.md': DocumentType.MD,
            '.jpg': DocumentType.IMAGE,
            '.jpeg': DocumentType.IMAGE,
            '.png': DocumentType.IMAGE,
        }
    
    def detect_type(self, filename: str) -> DocumentType:
        """检测文档类型"""
        ext = Path(filename).suffix.lower()
        return self.supported_formats.get(ext, DocumentType.UNKNOWN)
    
    def parse_pdf(self, file_path: str) -> str:
        """解析PDF文件"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n\n"
            return clean_extra_whitespace(text)
        except Exception as e:
            logger.error(f"PDF解析失败: {e}")
            raise
    
    def parse_docx(self, file_path: str) -> str:
        """解析DOCX文件"""
        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return clean_extra_whitespace(text)
        except Exception as e:
            logger.error(f"DOCX解析失败: {e}")
            raise
    
    def parse_txt(self, file_path: str) -> str:
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return clean_extra_whitespace(text)
        except Exception as e:
            logger.error(f"TXT解析失败: {e}")
            raise
    
    def parse_markdown(self, file_path: str) -> str:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return clean_extra_whitespace(text)
        except Exception as e:
            logger.error(f"Markdown解析失败: {e}")
            raise
    
    def parse_with_unstructured(self, file_path: str) -> str:
        """使用unstructured库解析文档（通用方法）"""
        try:
            elements = partition(filename=file_path)
            text = "\n".join([str(element) for element in elements])
            return clean_extra_whitespace(text)
        except Exception as e:
            logger.error(f"Unstructured解析失败: {e}")
            raise
    
    def parse(self, file_path: str, document_type: DocumentType) -> str:
        """根据文档类型调用相应的解析方法"""
        logger.info(f"开始解析文档: {file_path}, 类型: {document_type}")
        
        if document_type == DocumentType.PDF:
            return self.parse_pdf(file_path)
        elif document_type == DocumentType.DOCX:
            return self.parse_docx(file_path)
        elif document_type == DocumentType.TXT:
            return self.parse_txt(file_path)
        elif document_type == DocumentType.MD:
            return self.parse_markdown(file_path)
        elif document_type == DocumentType.IMAGE:
            # 图片文件需要OCR，这里先返回空
            logger.warning("图片文件需要OCR支持，当前版本暂不支持")
            return ""
        else:
            # 未知格式尝试使用unstructured
            try:
                return self.parse_with_unstructured(file_path)
            except:
                raise ValueError(f"不支持的文档格式: {document_type}")
    
    def extract_metadata(self, file_path: str, document_type: DocumentType) -> Dict[str, Any]:
        """提取文档元数据"""
        metadata = {
            "file_path": file_path,
            "file_size": os.path.getsize(file_path),
            "document_type": document_type.value,
            "parsed_with": "custom_parser"
        }
        
        try:
            if document_type == DocumentType.PDF:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata.update({
                        "page_count": len(pdf_reader.pages),
                        "pdf_metadata": pdf_reader.metadata
                    })
            elif document_type == DocumentType.DOCX:
                doc = DocxDocument(file_path)
                metadata.update({
                    "paragraph_count": len(doc.paragraphs)
                })
        except Exception as e:
            logger.warning(f"元数据提取失败: {e}")
        
        return metadata


class DocumentChunker:
    """文档分块器"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def chunk_text(self, text: str) -> List[str]:
        """将文本分块"""
        if not text:
            return []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            # 计算块结束位置
            end = start + self.chunk_size
            
            # 如果块结束位置超过文本长度，直接取到结尾
            if end >= text_length:
                chunk = text[start:text_length].strip()
                if chunk:
                    chunks.append(chunk)
                break
            
            # 尝试在句子边界处分割
            sentence_end = text.rfind('.', start, end)
            paragraph_end = text.rfind('\n\n', start, end)
            
            if paragraph_end > start and paragraph_end - start > self.chunk_size // 2:
                end = paragraph_end + 2  # 包括换行符
            elif sentence_end > start and sentence_end - start > self.chunk_size // 2:
                end = sentence_end + 1  # 包括句号
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            # 移动起始位置，考虑重叠
            start = end - self.chunk_overlap
        
        return chunks


# 单例实例
document_parser = DocumentParser()
document_chunker = DocumentChunker()